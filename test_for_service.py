import pandas as pd
import docx as dc
import random as rd

def get_answer(q_list, question):
    """Функция, которая находит заданный вопрос и выводит ответ из списка:
    :param q_list: файл со списком вопросов и ответов на них
    :param question: вопрос, ответ на который мы ищем"""
    value = False
    for element in q_list:
        if element['Q'] == question:
            value = str(element['A'])
            break
    return value

# задаем количество вопросов, которое необходимо отразить в бланке
n = int(input("Введите количество вопросов для тестирования: "))

#считываем информацию из таблицы с вопросами и ответами
q_data = pd.read_excel('QA.xlsx', sheet_name='data')
q_dict = q_data.to_dict(orient='records') #словарь вопрос-ответ из файла
q_list = [] #список только с вопросами из файла
test_pers = [] #в этот список будут занесены вопросы, которые выбрали random
#формируем список вопросов
for rec in q_dict:
    q_list.append(rec['Q'])

#считываем информацию из таблицы со списком сотрудников на тестирование
p_data = pd.read_excel('personal.xlsx', sheet_name='data', usecols='A')
p_list = p_data['Сотрудники для тестирования'].tolist() #список с сотрудниками

#для каждого сотрудника из списка формируем свой бланк
for pers in p_list:
    #выбираем N вопросов из списка q_list
    test_pers = rd.sample(q_list, n)
    #открываем шаблон, в который будем прописывать вопросы
    test = dc.Document('temp_test.docx')

    #блок, в котором заносим данные в таблицу шаблона (ФИО)
    test.tables #список со всеми таблицами в документе
    test.tables[0].cell(0, 1).text = pers

    #вносим выбранные вопросы в шаблон с тестом
    num = 1
    for element in test_pers:
        test.add_paragraph(str(num) + '. ' + str(element))
        test.add_paragraph('_' * 380)
        num += 1
    test.save('тест ' + pers + '.docx')
    print('Бланк с вопросами для сотрудника ' + pers + ' сформирован')
